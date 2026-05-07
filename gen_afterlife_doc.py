#!/usr/bin/env python3
"""Generate Afterlife technical study guide DOCX."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "/Users/unai/Desktop/Afterlife_Resumen_Tecnico.docx"

# ── Brand colours ──────────────────────────────────────────────────────────────
PURPLE   = RGBColor(0x6D, 0x28, 0xD9)   # deep purple for headings
ACCENT   = RGBColor(0xA8, 0x55, 0xF7)   # electric purple accent
DARK     = RGBColor(0x1E, 0x1B, 0x4B)   # near-black title
BODY     = RGBColor(0x1F, 0x29, 0x37)   # dark body text
SUBTLE   = RGBColor(0x4B, 0x55, 0x63)   # grey for subtitles/captions
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc: Document):
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "A855F7")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def style_run(run, bold=False, italic=False, size=None, color=None, font="Calibri"):
    run.font.name  = font
    run.font.bold  = bold
    run.font.italic = italic
    if size:  run.font.size  = Pt(size)
    if color: run.font.color.rgb = color

def add_title_block(doc: Document):
    """Create a styled title section."""
    # Coloured header table for visual impact
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1E1B4B")
    cell.width = Inches(6.5)

    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Title text
    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after  = Pt(6)
    r = p_title.add_run("AFTERLIFE")
    style_run(r, bold=True, size=32, color=WHITE)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("Resumen Técnico del Proyecto")
    style_run(r2, bold=False, size=18, color=RGBColor(0xC4, 0xB5, 0xFD))

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(6)
    r3 = p3.add_run("Guía de estudio para la presentación")
    style_run(r3, italic=True, size=13, color=RGBColor(0xE9, 0xD5, 0xFF))

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(20)
    r4 = p4.add_run("Unai Francés  ·  David Pereda")
    style_run(r4, bold=True, size=12, color=RGBColor(0xA5, 0xB4, 0xFC))

    doc.add_paragraph()  # spacer


def add_section_heading(doc: Document, number: int, title: str):
    """Add a numbered section heading with a coloured left border effect."""
    doc.add_paragraph()  # top spacer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)

    # Add left border in purple
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "6D28D9")
    pBdr.append(left)
    pPr.append(pBdr)

    r_num = p.add_run(f"{number}.  ")
    style_run(r_num, bold=True, size=14, color=ACCENT)
    r_title = p.add_run(title.upper())
    style_run(r_title, bold=True, size=14, color=PURPLE)

    add_horizontal_rule(doc)


def add_subheading(doc: Document, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(title)
    style_run(r, bold=True, size=12, color=PURPLE)
    return p


def add_body(doc: Document, text: str, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    style_run(r, size=11, color=BODY)
    return p


def add_bullet(doc: Document, text: str, bold_prefix: str = None):
    """Add a bullet point. If bold_prefix given, that part is bolded."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True, size=11, color=BODY)
        r2 = p.add_run(text)
        style_run(r2, size=11, color=BODY)
    else:
        r = p.add_run(text)
        style_run(r, size=11, color=BODY)
    return p


def add_code_inline(paragraph, text: str):
    """Add inline code-styled run."""
    r = paragraph.add_run(f"`{text}`")
    r.font.name  = "Courier New"
    r.font.size  = Pt(10)
    r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)


def add_qa_block(doc: Document, question: str, answer: str):
    """Add a Q&A styled block."""
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(8)
    p_q.paragraph_format.space_after  = Pt(3)
    r_q = p_q.add_run(f"P: {question}")
    style_run(r_q, bold=True, size=11, color=PURPLE)

    p_a = doc.add_paragraph()
    p_a.paragraph_format.space_after  = Pt(6)
    p_a.paragraph_format.left_indent  = Cm(0.5)
    r_a = p_a.add_run(f"R: {answer}")
    style_run(r_a, size=11, color=BODY)


def build_document():
    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # ── Default paragraph font ─────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── List Bullet style ──────────────────────────────────────────────────────
    try:
        lb = doc.styles["List Bullet"]
        lb.font.name = "Calibri"
        lb.font.size = Pt(11)
    except Exception:
        pass

    # ══════════════════════════════════════════════════════════════════════════
    # TITLE BLOCK
    # ══════════════════════════════════════════════════════════════════════════
    add_title_block(doc)

    # ── Intro paragraph ────────────────────────────────────────────────────────
    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(6)
    r_intro = intro_p.add_run(
        "Este documento es una guía de estudio técnica diseñada para que "
        "Unai Francés y David Pereda puedan repasar y explicar con claridad "
        "todos los aspectos del proyecto Afterlife durante su presentación. "
        "Cada sección cubre un componente clave de la arquitectura, con "
        "definiciones de los términos técnicos más importantes, ejemplos de "
        "código relevantes y explicaciones orientadas a la comunicación oral. "
        "Las secciones están ordenadas de lo general a lo específico."
    )
    style_run(r_intro, italic=True, size=11, color=SUBTLE)

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 1 — QUÉ ES AFTERLIFE
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 1, "¿Qué es Afterlife?")

    add_body(doc,
        "Afterlife es una aplicación móvil multiplataforma desarrollada con Flutter "
        "que permite a grupos de amigos organizar y vivir sus noches de fiesta de "
        "forma interactiva y competitiva. A continuación se detallan sus "
        "características principales:")

    bullets_sec1 = [
        ("Plataformas: ", "Android e iOS desde un único código base."),
        ("Crear noches: ", "Cualquier usuario puede crear una sesion de juego (una 'noche') con nombre personalizado y retos predefinidos o personalizados."),
        ("Unirse con código: ", "Los amigos se unen a la noche escaneando o introduciendo un código único."),
        ("Retos en tiempo real: ", "Los jugadores completan retos durante la noche; todos ven el progreso al instante."),
        ("Puntos y niveles: ", "Cada reto completado suma puntos. Los puntos acumulados a lo largo del tiempo aumentan el nivel del jugador."),
        ("Chat de grupo: ", "Chat en tiempo real dentro de cada noche y entre amigos."),
        ("Diario de salidas: ", "Registro histórico de todas las noches, fotos compartidas y estadísticas personales."),
        ("Logros: ", "Sistema de achievements que recompensa hitos (primera noche, 50 retos completados, etc.)."),
    ]
    for bold, rest in bullets_sec1:
        add_bullet(doc, rest, bold_prefix=bold)

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 2 — ARQUITECTURA GENERAL
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 2, "Arquitectura General")

    # Flutter + Dart
    add_subheading(doc, "Flutter + Dart")
    add_body(doc,
        "Flutter es el framework de código abierto de Google que permite crear "
        "aplicaciones nativas para Android, iOS y Web desde un único código fuente. "
        "Dart es el lenguaje de programación que usa Flutter: es compilado, "
        "fuertemente tipado y orientado a objetos.")

    flutter_bullets = [
        ("Framework: ", "Flutter compila a código nativo ARM/x64, lo que da un rendimiento equivalente al desarrollo nativo."),
        ("Widgets: ", "La unidad fundamental de Flutter. Todo en Flutter es un widget: botones, textos, imágenes, pantallas completas, incluso el padding."),
        ("StatelessWidget: ", "Widget sin estado interno. Su apariencia depende solo de los parámetros recibidos. Se reconstruye solo si el padre lo hace."),
        ("StatefulWidget: ", "Widget con estado mutable. Tiene un objeto State separado que sobrevive a las reconstrucciones. Usa setState() para notificar cambios."),
        ("Método build(): ", "Cada widget implementa build(BuildContext context), que devuelve el árbol de widgets a renderizar en pantalla."),
        ("Hot Reload: ", "Permite ver los cambios en el UI en tiempo real durante el desarrollo sin reiniciar la app."),
    ]
    for bold, rest in flutter_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # Provider
    add_subheading(doc, "Patrón Provider — Gestión de Estado")
    add_body(doc,
        "Provider es la solución de gestión de estado que usamos en Afterlife. "
        "Resuelve el problema de compartir datos (p. ej. el usuario autenticado) "
        "entre pantallas sin pasarlos manualmente por cada constructor.")

    provider_bullets = [
        ("ChangeNotifier: ", "Clase base que los providers extienden. Cuando los datos cambian, se llama a notifyListeners() y todos los widgets suscritos se reconstruyen."),
        ("ChangeNotifierProvider: ", "Widget que hace el provider disponible en todo el subárbol de widgets que contiene."),
        ("Consumer<T>: ", "Widget que se suscribe a un provider y se reconstruye solo cuando ese dato cambia. Más eficiente que reconstruir toda la pantalla."),
        ("context.watch<T>(): ", "Forma abreviada de suscribirse a un provider desde dentro de un widget."),
        ("context.read<T>(): ", "Lee el provider una sola vez sin suscribirse a cambios. Útil dentro de callbacks y métodos."),
        ("UserProvider: ", "Nuestro provider principal. Gestiona los datos del usuario autenticado, su perfil, nivel, logros y noche activa."),
        ("SettingsProvider: ", "Gestiona las preferencias de la app: tema visual (oscuro/claro), idioma y ajustes de notificaciones."),
    ]
    for bold, rest in provider_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # Estructura de carpetas
    add_subheading(doc, "Estructura de Carpetas del Proyecto")
    add_body(doc, "El proyecto sigue una arquitectura por funcionalidad (feature-based), "
        "organizando el código en capas claramente separadas:")

    folders = [
        ("lib/ ", "— Raíz de todo el código Dart de la aplicación."),
        ("lib/main.dart ", "— Punto de entrada. Inicializa Firebase y lanza la app con los providers configurados."),
        ("lib/services/ ", "— Lógica de negocio y comunicación con Firebase (night_service.dart, chat_service.dart, achievement_service.dart...)."),
        ("lib/providers/ ", "— Clases ChangeNotifier que gestionan el estado global (user_provider.dart, settings_provider.dart)."),
        ("lib/screens/ ", "— Pantallas secundarias de la app (perfil, historial, logros, etc.)."),
        ("lib/components/ ", "— Widgets reutilizables compartidos entre pantallas (tarjetas, botones, indicadores de carga...)."),
        ("lib/models/ ", "— Clases de datos que representan las entidades del dominio (NightModel, UserModel, AchievementModel...)."),
        ("lib/theme/ ", "— Definición centralizada de colores, tipografías y estilos del tema visual."),
    ]
    for bold, rest in folders:
        add_bullet(doc, rest, bold_prefix=bold)

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 3 — FIREBASE
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 3, "Firebase — El Backend Completo")

    add_body(doc,
        "Firebase es la plataforma Backend-as-a-Service (BaaS) de Google que usamos "
        "como infraestructura completa de Afterlife. Nos proporciona autenticación, "
        "base de datos, almacenamiento de archivos y más, sin necesidad de gestionar "
        "servidores propios.")

    # Auth
    add_subheading(doc, "Firebase Authentication")
    add_body(doc,
        "Gestiona el registro e inicio de sesión de usuarios de forma segura. "
        "Usamos el método de email + contraseña.")

    auth_bullets = [
        ("FirebaseAuth.instance: ", "Singleton global para interactuar con el servicio de autenticación."),
        ("authStateChanges(): ", "Devuelve un Stream<User?> que emite el usuario actual cada vez que cambia (login, logout, token refresh). Es la pieza central del flujo de autenticación."),
        ("StreamBuilder<User?>: ", "Widget Flutter que escucha el stream de autenticación y reconstruye el UI según el estado: mostrando LoginPage si no hay usuario o la app principal si está autenticado."),
        ("UID: ", "Identificador único universal que Firebase asigna a cada usuario al registrarse. Lo usamos como clave principal en todos los documentos de Firestore relacionados con ese usuario."),
        ("createUserWithEmailAndPassword(): ", "Método para registrar un nuevo usuario."),
        ("signInWithEmailAndPassword(): ", "Método para iniciar sesión."),
        ("signOut(): ", "Cierra la sesión del usuario actual y el stream emite null, llevando al usuario de vuelta a LoginPage automáticamente."),
    ]
    for bold, rest in auth_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # Firestore
    add_subheading(doc, "Cloud Firestore — Base de Datos")
    add_body(doc,
        "Cloud Firestore es la base de datos NoSQL en la nube de Firebase. "
        "Su principal característica es la sincronización en tiempo real: "
        "todos los clientes conectados ven los cambios al instante.")

    firestore_bullets = [
        ("Modelo de datos: ", "Colecciones → Documentos → Campos. Una Colección contiene Documentos; cada Documento tiene Campos (pares clave-valor)."),
        ("Colecciones principales: ", "users, nights, achievements, moments, chats."),
        ("Lectura única: ", ".get() — obtiene el documento una sola vez."),
        ("Stream en tiempo real: ", ".snapshots() — devuelve un Stream que emite el documento cada vez que se modifica en el servidor."),
        ("Crear documento: ", "_firestore.collection('nights').add(data) — genera un ID automático."),
        ("Actualizar documento: ", "_firestore.collection('nights').doc(id).update(data) — actualiza solo los campos indicados."),
        ("FieldValue.arrayUnion([x]): ", "Añade un elemento a un array en Firestore sin necesidad de leer el array completo primero ni riesgo de sobreescribirlo."),
        ("FieldValue.serverTimestamp(): ", "Inserta el timestamp del servidor (no del dispositivo), garantizando coherencia entre zonas horarias."),
        ("Transacciones (runTransaction): ", "Operación atómica que lee y escribe documentos de forma segura, sin que ningún otro cliente pueda modificarlos en medio de la operación. Las usamos al unirse a una noche."),
        ("Batch writes: ", "Agrupan múltiples operaciones de escritura en una sola petición. Lo usamos para crear los 10 logros iniciales de golpe."),
    ]
    for bold, rest in firestore_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # Storage
    add_subheading(doc, "Firebase Storage — Almacenamiento de Archivos")
    add_body(doc,
        "Firebase Storage almacena archivos binarios (imágenes, vídeos) en la nube. "
        "Separar los archivos de la base de datos es fundamental porque Firestore "
        "tiene un límite de 1 MB por documento.")

    storage_bullets = [
        ("Ruta de fotos: ", "nights/{nightId}/photos/{timestamp}.jpg — estructura organizada por noche."),
        ("ref.putData(bytes): ", "Sube los bytes de la imagen a Firebase Storage."),
        ("ref.getDownloadURL(): ", "Una vez subida, obtiene la URL pública de descarga."),
        ("Guardar URL en Firestore: ", "Solo se almacena la URL (texto) en el documento de la noche, no los bytes. Esto elimina el límite práctico de fotos."),
        ("Image.network(url): ", "Widget Flutter para mostrar imágenes desde una URL HTTP."),
    ]
    for bold, rest in storage_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 4 — FLUJO main.dart
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 4, "Flujo de la Aplicación — main.dart")

    add_subheading(doc, "Inicialización de la App")
    add_body(doc,
        "El archivo main.dart es el punto de entrada de cualquier aplicación Flutter. "
        "En Afterlife realiza tres tareas críticas antes de mostrar ninguna pantalla:")

    init_bullets = [
        ("main() async: ", "La función main es asíncrona porque necesita await para esperar a que Firebase esté listo."),
        ("WidgetsFlutterBinding.ensureInitialized(): ", "Debe llamarse antes de usar cualquier plugin de Flutter. Inicializa el binding entre el framework Dart y el motor nativo."),
        ("Firebase.initializeApp(): ", "Conecta la app con el proyecto Firebase usando el archivo de configuración (google-services.json en Android, GoogleService-Info.plist en iOS)."),
        ("MultiProvider: ", "Configura todos los providers (UserProvider, SettingsProvider) en la raíz del árbol de widgets para que estén disponibles en toda la app."),
    ]
    for bold, rest in init_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    add_subheading(doc, "OnboardingWrapper — El Enrutador Inteligente")
    add_body(doc,
        "OnboardingWrapper es el widget raíz que decide qué pantalla mostrar al usuario "
        "según su estado. Es el cerebro de la navegación inicial:")

    wrapper_bullets = [
        ("StreamBuilder<User?>: ", "Escucha authStateChanges() de Firebase Authentication."),
        ("Estado cargando: ", "Mientras el stream no ha emitido su primer valor, se muestra un spinner de carga."),
        ("Sin usuario (null): ", "El stream emite null → se muestra LoginPage."),
        ("Usuario sin tutorial: ", "El stream emite un usuario, pero en sus datos de Firestore hasSeenTutorial == false → se muestra TutorialScreen."),
        ("Usuario con tutorial visto: ", "El stream emite un usuario y hasSeenTutorial == true → se muestra MainScreen (la app completa)."),
    ]
    for bold, rest in wrapper_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    add_subheading(doc, "Deep Links (Versión Web)")
    add_body(doc,
        "En la versión web de Afterlife, cuando alguien comparte un enlace del tipo "
        "https://afterlife.app?nightId=ABC123, la app detecta el parámetro nightId "
        "en la URL al arrancar y navega directamente a esa noche, permitiendo unirse "
        "sin buscar manualmente.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 5 — SISTEMA DE NOCHES
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 5, "Sistema de Noches — night_service.dart")

    add_body(doc,
        "NightService centraliza toda la lógica de creación, gestión y seguimiento "
        "de las noches de juego. Es el servicio más complejo de la app.")

    add_subheading(doc, "Ciclo de Vida de una Noche")
    lifecycle_bullets = [
        ("waiting: ", "Estado inicial. La noche ha sido creada pero aún no ha empezado. Los jugadores se unen en este estado."),
        ("active: ", "La noche está en curso. Los jugadores pueden completar retos y chatear."),
        ("finished: ", "La noche ha terminado. Las estadísticas se calculan y guardan definitivamente."),
    ]
    for bold, rest in lifecycle_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    add_subheading(doc, "Crear una Noche")
    add_body(doc,
        "Al crear una noche se genera un documento en la colección nights con los "
        "siguientes campos principales: nombre de la noche, UID del anfitrión (host), "
        "lista de jugadores (inicialmente solo el anfitrión), lista de retos con su "
        "estado (completado/pendiente), estado general (waiting), código de invitación "
        "único, timestamps de creación, y arrays vacíos para fotos y mensajes de chat.")

    add_subheading(doc, "Unirse a una Noche — Transacción Firestore")
    add_body(doc,
        "El proceso de unirse a una noche usa una transacción para evitar la race "
        "condition donde dos usuarios intentan unirse simultáneamente a una noche "
        "que solo tiene una plaza libre:")

    join_bullets = [
        ("Paso 1 — Leer: ", "Dentro de la transacción, se lee el documento actual de la noche."),
        ("Paso 2 — Validar: ", "Se comprueba si el usuario ya está en la lista, si la noche está llena o si ya ha empezado."),
        ("Paso 3 — Escribir: ", "Solo si todas las validaciones pasan, se añade el jugador con FieldValue.arrayUnion()."),
        ("Atomicidad: ", "Si entre el paso 1 y el paso 3 otro usuario modificó el documento, Firestore cancela y reintenta la transacción automáticamente."),
    ]
    for bold, rest in join_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    add_subheading(doc, "Completar un Reto")
    add_body(doc,
        "Al marcar un reto como completado, NightService actualiza en una sola "
        "transacción: el estado del reto (completed = true), quién lo completó (UID), "
        "el timestamp de completado, y los puntos del jugador en el documento de la noche.")

    add_subheading(doc, "Streaming en Tiempo Real")
    add_body(doc,
        "streamNight(nightId) devuelve un Stream<DocumentSnapshot> que todos los "
        "jugadores de la noche escuchan simultáneamente. Cada vez que cualquier "
        "jugador completa un reto, añade una foto o escribe un mensaje, Firestore "
        "emite una actualización y el UI de todos los participantes se reconstruye "
        "automáticamente a través del StreamBuilder correspondiente.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 6 — GESTIÓN DE ESTADO (UserProvider)
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 6, "Gestión de Estado — UserProvider")

    add_body(doc,
        "UserProvider es el eje central de la gestión de estado de Afterlife. "
        "Extiende ChangeNotifier y actúa como puente entre Firebase Authentication, "
        "Firestore y el UI de la app.")

    up_bullets = [
        ("authStateChanges().listen(): ", "UserProvider se suscribe al stream de autenticación al iniciarse. Cuando hay un login, carga automáticamente los datos del usuario desde Firestore."),
        ("userData: ", "Mapa con todos los datos del perfil del usuario (nombre, avatar, nivel, puntos, estadísticas)."),
        ("isLoading: ", "Booleano que indica si se están cargando datos. Los widgets lo usan para mostrar spinners."),
        ("activeNightId: ", "ID de la noche en la que el usuario está participando actualmente (si la hay)."),
        ("unlockedAchievements: ", "Lista de logros desbloqueados por el usuario."),
        ("OfflineService: ", "Si la carga de datos falla por falta de red, UserProvider usa datos en caché local para que la app siga funcionando offline."),
    ]
    for bold, rest in up_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    add_subheading(doc, "Actualización de Estadísticas Post-Noche")
    add_body(doc,
        "Al finalizar una noche, updateAfterNight() realiza todas las actualizaciones "
        "de perfil del usuario en Firestore:")

    stats_bullets = [
        ("Puntos: ", "Suma los puntos ganados durante la noche al total acumulado."),
        ("Noches completadas: ", "Incrementa el contador de noches jugadas."),
        ("Retos totales: ", "Suma los retos completados en esta noche."),
        ("Nivel: ", "Recalcula el nivel usando LevelCalculator.calculate(totalPoints)."),
        ("Racha (streak): ", "Comprueba si la salida es en un día consecutivo al de la última noche. Si sí, incrementa la racha; si no, la resetea a 1."),
    ]
    for bold, rest in stats_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 7 — LOGROS
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 7, "Sistema de Logros — AchievementService")

    add_body(doc,
        "El sistema de logros (achievements) gamifica la experiencia de uso a largo "
        "plazo, recompensando a los usuarios más activos con insignias y puntos extra.")

    ach_bullets = [
        ("Colección achievements: ", "Los logros se definen y almacenan en Firestore. Al arrancar la app, si la colección está vacía se crean los 10 logros por defecto usando un batch write."),
        ("checkAndUnlockAchievements(): ", "Función llamada después de cada noche. Compara las estadísticas del usuario (noches jugadas, retos completados, puntos totales) contra los requisitos de cada logro."),
        ("Desbloqueo: ", "Si el usuario cumple el requisito de un logro que aún no tiene, se guarda en su perfil con timestamp y puntos ganados, y se lanza una notificación en pantalla."),
        ("Logro 'NO VETERANO': ", "Primera noche completada."),
        ("Logro 'LEYENDA': ", "50 retos completados en total."),
        ("Logro 'NOCTÁMBULO': ", "100 noches jugadas."),
    ]
    for bold, rest in ach_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 8 — CHAT
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 8, "Sistema de Chat")

    add_subheading(doc, "Chat entre Amigos — ChatService")
    add_body(doc,
        "El chat privado entre dos usuarios funciona de la siguiente manera:")

    chat_bullets = [
        ("ID de conversación: ", "Se genera combinando los UIDs de los dos participantes ordenados alfabéticamente y separados por underscore. Esto garantiza que siempre sea el mismo ID independientemente de quién inicie la conversación."),
        ("Estructura Firestore: ", "chats/{chatId}/messages/{messageId} — cada mensaje es un documento dentro de la subcolección messages del chat."),
        ("Tiempo real: ", "Los mensajes se escuchan con .snapshots() ordenados por timestamp, de forma que el chat se actualiza instantáneamente."),
        ("Límite de caracteres: ", "500 caracteres por mensaje, validado en cliente antes de enviar."),
    ]
    for bold, rest in chat_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    add_subheading(doc, "Chat de Grupo — NightChatService")
    add_body(doc,
        "El chat de grupo durante una noche activa tiene características específicas:")

    night_chat_bullets = [
        ("Ruta: ", "nights/{nightId}/chat/{messageId} — los mensajes del grupo son una subcolección del documento de la noche."),
        ("Acceso restringido: ", "Solo los usuarios cuyo UID aparezca en el array de jugadores de la noche pueden escribir mensajes (validado mediante reglas de seguridad de Firestore)."),
        ("Mismo patrón: ", "Igual que el chat privado, usa .snapshots() para tiempo real y muestra el nombre y avatar del remitente en cada mensaje."),
    ]
    for bold, rest in night_chat_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 9 — NAVEGACIÓN
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 9, "Navegación en Flutter")

    add_body(doc,
        "Flutter gestiona la navegación mediante una pila (stack) de rutas. "
        "Entender este sistema es importante para explicar cómo el usuario se mueve "
        "entre pantallas:")

    nav_bullets = [
        ("Navigator.push(): ", "Añade una nueva pantalla encima de la pila. El usuario puede volver atrás con el botón o gesto."),
        ("Navigator.pop(): ", "Elimina la pantalla actual de la pila, volviendo a la anterior."),
        ("Navigator.pushReplacement(): ", "Reemplaza la pantalla actual por una nueva. No se puede volver atrás. Usado en login → home."),
        ("Rutas con nombre: ", "Definidas en MaterialApp.routes como un mapa de String → WidgetBuilder. Ejemplo: '/edit-profile', '/settings'. Permiten navegar sin importar la clase."),
        ("PageView + PageController: ", "Componente que permite mostrar varias páginas deslizables (para el nav bar inferior con las pestañas principales)."),
        ("NeverScrollableScrollPhysics(): ", "Desactiva el deslizamiento táctil en el PageView para que solo se pueda cambiar de pestaña pulsando el nav bar, evitando transiciones accidentales."),
    ]
    for bold, rest in nav_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 10 — UI Y TEMA VISUAL
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 10, "UI y Tema Visual")

    add_subheading(doc, "Sistema de Temas")
    add_body(doc,
        "Afterlife soporta tema oscuro y tema claro, gestionados de forma centralizada:")

    theme_bullets = [
        ("AfterlifeTheme: ", "Clase que define tanto darkTheme como lightTheme usando ThemeData de Flutter."),
        ("SettingsProvider: ", "Guarda la preferencia del usuario en almacenamiento local y la aplica pasándola a themeMode en MaterialApp."),
        ("Colores principales: ", "Morado eléctrico #A855F7, rosa neón #EC4899, verde ácido #84CC16, azul cyan #06B6D4. Definen la identidad visual nocturna y energética de la app."),
    ]
    for bold, rest in theme_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    add_subheading(doc, "Widgets Clave del UI")

    ui_bullets = [
        ("StreamBuilder: ", "Widget que escucha un Stream y reconstruye su subtree cada vez que llega un nuevo valor. Fundamental para el tiempo real en Afterlife."),
        ("Consumer<T>: ", "Lee un provider y reconstruye solo su subtree cuando ese provider notifica cambios. Más eficiente que usar context.watch() en el widget raíz de una pantalla."),
        ("AnimatedContainer: ", "Versión animada de Container. Cuando sus propiedades (tamaño, color, padding) cambian, la transición es suave y automática."),
        ("ShaderMask: ", "Aplica un gradiente de colores al widget hijo. Usado para crear los textos con gradiente de morado a rosa que caracterizan el estilo de la app."),
        ("PageView: ", "Permite deslizarse entre páginas. Usado tanto en el tutorial de bienvenida como en la navegación principal por pestañas."),
        ("ModalBottomSheet: ", "Muestra una bandeja desde la parte inferior de la pantalla. Usado para chats, ajustes rápidos y confirmaciones sin abandonar la pantalla actual."),
    ]
    for bold, rest in ui_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 11 — FOTOS Y MULTIMEDIA
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 11, "Fotos y Multimedia — Evolución Técnica")

    add_subheading(doc, "Problema Original (v1)")
    add_body(doc,
        "En la primera versión de la app, las fotos se guardaban como bytes directamente "
        "en documentos de Firestore. Esto generaba dos problemas graves:")

    v1_bullets = [
        ("Límite de tamaño: ", "Firestore tiene un límite de 1 MB por documento. Las imágenes ocupaban la mayor parte de ese espacio."),
        ("Máximo de fotos: ", "Solo se podían guardar 5 fotos por noche, y además con una compresión agresiva (máx. 300 KB cada una)."),
        ("Rendimiento: ", "Leer un documento con bytes embebidos es lento y costoso en términos de ancho de banda."),
    ]
    for bold, rest in v1_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    add_subheading(doc, "Solución Actual con Firebase Storage (v2)")
    add_body(doc,
        "La migración a Firebase Storage resolvió todos estos problemas:")

    v2_bullets = [
        ("Subida: ", "ref.putData(bytes) — sube la imagen a Firebase Storage, que no tiene los límites de Firestore."),
        ("URL pública: ", "ref.getDownloadURL() — devuelve una URL HTTPS pública para la imagen subida."),
        ("Guardar en Firestore: ", "Solo se guarda la URL (string corto) en el documento de la noche usando FieldValue.arrayUnion([url])."),
        ("Mostrar en Flutter: ", "Image.network(url) carga y cachea la imagen desde la URL."),
        ("Sin límite práctico: ", "Ahora se pueden guardar decenas de fotos por noche sin restricción de tamaño en Firestore."),
    ]
    for bold, rest in v2_bullets:
        add_bullet(doc, rest, bold_prefix=bold)

    # ══════════════════════════════════════════════════════════════════════════
    # SECCIÓN 12 — PREGUNTAS FRECUENTES
    # ══════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 12, "Preguntas Frecuentes — Preparación para la Presentación")

    add_body(doc,
        "Esta sección recoge las preguntas técnicas más probables durante la presentación, "
        "con respuestas concisas y directas listas para verbalizar:")

    add_qa_block(doc,
        "¿Por qué Flutter y no React Native o desarrollo nativo?",
        "Flutter compila directamente a código nativo ARM/x64 para Android e iOS desde "
        "un único código base, sin el puente JavaScript de React Native. Esto significa "
        "mejor rendimiento, especialmente en animaciones. Dart se compila AOT (Ahead Of Time), "
        "lo que da tiempos de arranque muy rápidos. Además, el sistema de widgets de Flutter "
        "dibuja directamente en la pantalla con su propio motor de renderizado (Skia/Impeller), "
        "dando consistencia visual total entre plataformas.")

    add_qa_block(doc,
        "¿Por qué Firebase?",
        "Firebase ofrece autenticación, base de datos en tiempo real, almacenamiento de archivos "
        "y hosting bajo un mismo SDK y proyecto. Para Afterlife, que necesita sincronización "
        "en tiempo real entre múltiples usuarios, sería muy costoso implementar un backend "
        "propio desde cero. Firebase nos permitió centrarnos en la experiencia de usuario "
        "y la lógica del juego. También escala automáticamente.")

    add_qa_block(doc,
        "¿Qué es una transacción en Firestore y por qué la usamos?",
        "Una transacción es una operación atómica: lee uno o más documentos y luego escribe "
        "en ellos, garantizando que ningún otro cliente puede modificarlos entre la lectura "
        "y la escritura. En Afterlife la usamos al unirse a una noche: si dos usuarios intentan "
        "unirse al mismo tiempo a una noche con una sola plaza libre, sin transacción ambos "
        "podrían leer que hay hueco y añadirse, superando el límite. Con la transacción, "
        "solo uno lo consigue y el otro recibe un error. Esto se llama race condition y la "
        "transacción la previene.")

    add_qa_block(doc,
        "¿Qué es un Stream en Flutter/Firestore?",
        "Un Stream es un flujo de datos asíncrono que emite valores a lo largo del tiempo, "
        "como un canal de TV que sigue transmitiendo en lugar de un video que descargas una "
        "vez. En Firestore, llamar a .snapshots() sobre una colección o documento devuelve "
        "un Stream que emite el estado actual del dato y luego emite una nueva versión cada "
        "vez que alguien lo modifica en la base de datos. Con StreamBuilder en Flutter, "
        "el widget se reconstruye automáticamente con cada nueva emisión, creando así la "
        "sincronización en tiempo real.")

    add_qa_block(doc,
        "¿Qué es Provider y por qué lo usamos?",
        "Provider es el sistema de gestión de estado recomendado por el equipo de Flutter. "
        "Sin él, para pasar el objeto del usuario actual desde la pantalla raíz hasta un "
        "widget profundamente anidado habría que pasarlo como parámetro en cada constructor "
        "intermedio (prop drilling). Con Provider, el dato se coloca en el árbol una sola "
        "vez (con ChangeNotifierProvider) y cualquier widget puede acceder a él directamente "
        "con context.watch() o Consumer<T>. Cuando el dato cambia, solo los widgets suscritos "
        "se reconstruyen, no toda la pantalla.")

    add_qa_block(doc,
        "¿Cómo funciona el sistema de niveles?",
        "LevelCalculator.calculate(totalPoints) es una función matemática que toma los puntos "
        "totales acumulados del usuario y devuelve su nivel actual. Funciona con umbrales "
        "progresivos: los primeros niveles requieren pocos puntos pero los niveles altos "
        "requieren exponencialmente más. Los puntos se ganan completando retos durante las "
        "noches y desbloqueando logros. El nivel se recalcula y actualiza en Firestore "
        "cada vez que el usuario termina una noche.")

    add_qa_block(doc,
        "¿Cómo se garantiza la seguridad de los datos?",
        "Firebase Security Rules permiten definir quién puede leer y escribir cada documento. "
        "En Afterlife, un usuario solo puede leer o modificar su propio documento en la "
        "colección users. Solo los participantes de una noche pueden leer los datos de esa "
        "noche o escribir mensajes en su chat. Estas reglas se evalúan en el servidor de "
        "Firebase, por lo que no pueden ser saltadas desde el cliente aunque el código "
        "de la app fuese modificado.")

    add_qa_block(doc,
        "¿Qué diferencia hay entre StatelessWidget y StatefulWidget?",
        "Un StatelessWidget no tiene estado interno: su apariencia depende únicamente de "
        "los parámetros que recibe y nunca cambia por sí solo. Un StatefulWidget tiene un "
        "objeto State asociado que puede cambiar, y cuando cambia (llamando a setState()), "
        "Flutter reconstruye solo ese widget. En Afterlife usamos StatelessWidgets para "
        "componentes puramente visuales (tarjetas, botones) y StatefulWidgets para elementos "
        "con interacción o datos locales (formularios, animaciones locales).")

    # ── Final spacer ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run("— Fin del documento —")
    style_run(r_end, italic=True, size=10, color=SUBTLE)

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
